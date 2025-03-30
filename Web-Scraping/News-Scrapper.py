from selenium import webdriver         
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Pt
import pandas as pd

chrome_options = Options()
chrome_options.add_argument('--headless')
chrome_options.add_argument('--disable-gpu')
chrome_options.add_argument('--no-sandbox')
driver = webdriver.Chrome(options=chrome_options)

urls = ['https://www.educationtimes.com/article/newsroom/99734828/cbse-releases-list-of-fake-social-media-handles-to-avoid-misinformation-more-details-here' ,
        'https://www.educationtimes.com/article/newsroom/99734829/upsc-cse-2024-notification-releases-today;-check-eligibility-expected-exam-dates' ,
        'https://www.educationtimes.com/article/newsroom/99734831/dsssb-2024-recruitment-application-process-for-567-multi-tasking-staff-posts-begins-find-details-here' ,
        'https://www.educationtimes.com/article/newsroom/99734832/british-council-announces-great-scholarships-2024-for-pg-courses-in-science-technology-law-and-humanities',
        'https://www.educationtimes.com/article/newsroom/99734834/nzea-scholarship-for-27-indian-students-who-will-study-in-globally-ranked-nz-universities',
        'https://www.educationtimes.com/article/newsroom/99734838/board-exams-2024-cbse-class-x-xii-exams-begin-today-check-important-guidelines-here',
        'https://www.educationtimes.com/article/newsroom/99734839/nua-o-scholarship-odisha-launches-financial-assistance-scheme-for-ug-pg-students-details-inside',
        'https://www.educationtimes.com/article/newsroom/99734841/ipmat-2024-iim-indore-begins-registrations;-check-details-here' ,
        'https://www.educationtimes.com/article/newsroom/99734843/isro-ursc-recruitment-2024-is-underway-for-224-posts-more-details-here' ,
        'https://www.educationtimes.com/article/newsroom/99734844/sail-2024-recruitment-admit-card-for-exam-relating-to-hiring-for-technical-posts-released-find-details-here' ,
        'https://www.educationtimes.com/article/newsroom/99734845/ignou-to-close-the-registration-window-for-the-january-semester-today-check-details-here']

# Create a DataFrame to store the data
df = pd.DataFrame(columns=['Date','Title', 'Headline'])

for url in urls:
    driver.get(url)

    # Find the title of the news
    title_xpath = '//*[@id="__next"]/div[4]/div[3]/div/div/div[1]/section/div[1]/div[1]/h1'
    news_title = driver.find_element(By.XPATH, title_xpath).text

    # Find elements for headlines
    headline_xpath = '//*[@id="__next"]/div[4]/div[3]/div/div/div[1]/section/div[1]/div[3]/dl/div'
    headline_elements = driver.find_elements(By.XPATH, headline_xpath)

    # Create a list to store the data for each URL
    data = []

    # Extract and print the text from each headline along with the title
    for headline_element in headline_elements:
        print(f"Title: {news_title}")
        print(headline_element.text)
        print()

        # Append the data to the DataFrame
        data.append({'Title': news_title, 'Headline': headline_element.text})

    # Append the data for each URL to the main DataFrame
    df = pd.concat([df, pd.DataFrame(data)])

# Export the DataFrame to a Word document using python-docx
doc_path = '/content/drive/MyDrive/scrap_news14_15feb.docx'
doc = Document()

# Add DataFrame content to Word document
for index, row in df.iterrows():
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(row['Title'])
    title_run.bold = True
    title_run.font.size = Pt(16)  # You can adjust the font size as needed
    doc.add_paragraph (row['Headline'])
    doc.add_paragraph()  # Add an empty line between entries

# Save the Word document
doc.save(doc_path)

# Close the browser window
driver.quit()

print(f"Data has been exported to {doc_path}")

